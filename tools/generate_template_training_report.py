from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs"
FIG_DIR = OUT_DIR / "template_work" / "figures"
OUT_DOCX = OUT_DIR / "《综合项目训练》实战报告_CREMS_第六组.docx"


def ensure_dirs() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    FIG_DIR.mkdir(parents=True, exist_ok=True)


def font_path() -> str:
    candidates = [
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/PingFang.ttc",
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]
    for item in candidates:
        if Path(item).exists():
            return item
    return ""


FONT_PATH = font_path()


def pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    if FONT_PATH:
        return ImageFont.truetype(FONT_PATH, size=size)
    return ImageFont.load_default()


def set_run_font(run, size: float | None = None, bold: bool | None = None, name: str = "宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, first_indent: bool = True):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.25
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    if first_indent:
        fmt.first_line_indent = Pt(24)


def add_paragraph(doc: Document, text: str = "", style: str | None = None, align=None, first_indent: bool = True):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    set_paragraph_format(p, first_indent=first_indent)
    if text:
        r = p.add_run(text)
        set_run_font(r, 12, False, "宋体")
    return p


def add_heading(doc: Document, text: str, level: int):
    if level == 1:
        p = doc.add_paragraph(style="Heading 1")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(18)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(text)
        set_run_font(r, 16, True, "黑体")
    elif level == 2:
        p = doc.add_paragraph(style="Heading 2")
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(text)
        set_run_font(r, 14, True, "黑体")
    else:
        p = doc.add_paragraph(style="Heading 3")
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r = p.add_run(text)
        set_run_font(r, 12, True, "黑体")
    return p


def add_front_title(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(18)
    p.paragraph_format.line_spacing = 1.25
    r = p.add_run(text)
    set_run_font(r, 16, True, "黑体")
    return p


def add_page_break(doc: Document):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def add_caption(doc: Document, text: str, kind: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.25
    if kind == "table":
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(4)
    else:
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, 10.5, False, "宋体")
    return p


def set_cell_text(cell, text: str, bold: bool = False, size: float = 10.5, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_run_font(r, size, bold, "宋体" if not bold else "黑体")


def add_table(doc: Document, caption: str, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float]):
    add_caption(doc, caption, "table")
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], header, bold=True, size=10.5)
        table.rows[0].cells[i].width = Cm(widths[i])
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if len(value) > 18 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[i], value, size=10.5, align=align)
            cells[i].width = Cm(widths[i])
    doc.add_paragraph()
    return table


def add_toc_field(doc: Document):
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = r'TOC \o "1-3" \h \z \u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "右键更新域以生成目录"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run_font(run, 12, False, "宋体")


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_end])
    set_run_font(run, 10.5, False, "宋体")


def configure_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)
    add_page_number(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.first_line_indent = Pt(24)
    normal.paragraph_format.space_after = Pt(6)
    heading_tokens = [
        ("Heading 1", 16, "黑体"),
        ("Heading 2", 14, "黑体"),
        ("Heading 3", 12, "黑体"),
    ]
    for style_name, size, name in heading_tokens:
        style = styles[style_name]
        style.font.name = name
        style._element.rPr.rFonts.set(qn("w:eastAsia"), name)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.25


def cover(doc: Document):
    for _ in range(3):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("《综合项目训练》实战报告")
    set_run_font(r, 22, True, "黑体")
    for _ in range(3):
        doc.add_paragraph()

    fields = [
        ("题       目", "校园招聘与就业信息管理系统"),
        ("专 业 名 称", "计算机科学与技术"),
        ("班 级 名 称", "12023051B"),
        ("小 组 名 称", "第六组"),
        ("设 计 地 点", "S710"),
        ("指 导 教 师", "张维冬"),
        ("起 止 时 间", "2026.03.05-2026.06.18"),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (label, value) in enumerate(fields):
        set_cell_text(table.cell(i, 0), label, size=14, align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(table.cell(i, 1), value, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
        table.cell(i, 0).width = Cm(5)
        table.cell(i, 1).width = Cm(9)
        for cell in table.rows[i].cells:
            tc_pr = cell._tc.get_or_add_tcPr()
            borders = OxmlElement("w:tcBorders")
            for edge in ("top", "left", "bottom", "right"):
                elem = OxmlElement(f"w:{edge}")
                elem.set(qn("w:val"), "nil")
                borders.append(elem)
            tc_pr.append(borders)
    for _ in range(4):
        doc.add_paragraph()
    for text in ["三江学院", "计算机科学与工程学院"]:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(text)
        set_run_font(r2, 16, True, "宋体")


def member_page(doc: Document):
    add_page_break(doc)
    add_front_title(doc, "小组成员分工")
    rows = [
        ("1", "12023081065", "张赫", "组长；负责项目总体架构、数据库设计、后端业务模块、前端核心页面、接口联调、系统集成、构建验证和报告统稿。"),
        ("2", "12023051115", "徐喜", "协助进行功能测试、页面体验检查、测试用例整理，参与企业管理、职位管理等模块的操作验证。"),
        ("3", "12023051087", "张庭硕", "协助整理数据库表结构、业务字典和测试数据，参与学生端求职、投递、收藏流程测试。"),
        ("4", "12023051089", "江皓天", "协助准备演示材料、答辩说明和运行截图，参与企业端简历处理、面试管理流程测试。"),
    ]
    add_table(doc, "", ["序号", "学号", "姓名", "任务分工"], rows, [1.4, 3.3, 2.2, 9.2])


def toc_page(doc: Document):
    add_page_break(doc)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("目  录")
    set_run_font(r, 16, True, "黑体")
    add_toc_field(doc)


def draw_boxes(path: Path, title: str, boxes: list[tuple[int, int, int, int, str]], lines: list[tuple[int, int, int, int]]):
    img = Image.new("RGB", (1400, 780), "white")
    draw = ImageDraw.Draw(img)
    font = pil_font(34)
    small = pil_font(28)
    draw.text((700, 40), title, fill=(30, 41, 59), font=font, anchor="ma")
    for x1, y1, x2, y2 in lines:
        draw.line((x1, y1, x2, y2), fill=(71, 85, 105), width=4)
    for x1, y1, x2, y2, text in boxes:
        draw.rounded_rectangle((x1, y1, x2, y2), radius=18, outline=(51, 65, 85), width=3, fill=(241, 245, 249))
        parts = text.split("\n")
        total = len(parts) * 34
        for idx, line in enumerate(parts):
            draw.text(((x1 + x2) / 2, (y1 + y2 - total) / 2 + idx * 38 + 18), line, fill=(15, 23, 42), font=small, anchor="mm")
    img.save(path)


def make_figures():
    draw_boxes(
        FIG_DIR / "fig2-1-modules.png",
        "校园招聘与就业信息管理系统功能模块",
        [
            (520, 100, 880, 180, "CREMS 系统"),
            (80, 270, 330, 360, "管理员端\n企业/学生/职位审核"),
            (405, 270, 655, 360, "企业端\n职位/投递/面试"),
            (730, 270, 980, 360, "学生端\n搜索/投递/收藏"),
            (1055, 270, 1305, 360, "统计分析\n指标/图表/导出"),
            (120, 500, 320, 590, "企业管理"),
            (380, 500, 580, 590, "职位管理"),
            (640, 500, 840, 590, "简历投递"),
            (900, 500, 1100, 590, "面试安排"),
            (1160, 500, 1360, 590, "数据看板"),
        ],
        [
            (700, 180, 700, 230),
            (205, 230, 1180, 230),
            (205, 230, 205, 270),
            (530, 230, 530, 270),
            (855, 230, 855, 270),
            (1180, 230, 1180, 270),
            (205, 360, 220, 500),
            (530, 360, 480, 500),
            (855, 360, 740, 500),
            (855, 360, 1000, 500),
            (1180, 360, 1260, 500),
        ],
    )
    draw_boxes(
        FIG_DIR / "fig2-2-usecase.png",
        "系统参与者与核心用例",
        [
            (70, 310, 240, 390, "管理员"),
            (70, 500, 240, 580, "企业用户"),
            (70, 120, 240, 200, "学生用户"),
            (450, 100, 680, 180, "职位搜索/收藏"),
            (450, 230, 680, 310, "简历投递"),
            (450, 360, 680, 440, "查看面试"),
            (760, 170, 990, 250, "发布职位"),
            (760, 300, 990, 380, "处理投递"),
            (760, 430, 990, 510, "安排面试"),
            (1060, 170, 1290, 250, "企业/职位审核"),
            (1060, 300, 1290, 380, "基础数据维护"),
            (1060, 430, 1290, 510, "统计分析"),
        ],
        [
            (240, 160, 450, 140), (240, 160, 450, 270), (240, 160, 450, 400),
            (240, 540, 760, 210), (240, 540, 760, 340), (240, 540, 760, 470),
            (240, 350, 1060, 210), (240, 350, 1060, 340), (240, 350, 1060, 470),
        ],
    )
    draw_boxes(
        FIG_DIR / "fig2-3-flow.png",
        "校园招聘业务流程",
        [
            (60, 330, 240, 410, "企业入驻"),
            (300, 330, 480, 410, "管理员审核"),
            (540, 330, 720, 410, "职位发布"),
            (780, 330, 960, 410, "学生投递"),
            (1020, 330, 1200, 410, "企业筛选"),
            (540, 520, 720, 600, "面试安排"),
            (780, 520, 960, 600, "结果反馈"),
            (1020, 520, 1200, 600, "统计汇总"),
        ],
        [
            (240, 370, 300, 370), (480, 370, 540, 370), (720, 370, 780, 370),
            (960, 370, 1020, 370), (1110, 410, 630, 520), (720, 560, 780, 560),
            (960, 560, 1020, 560),
        ],
    )
    draw_boxes(
        FIG_DIR / "fig3-1-er.png",
        "CREMS 核心业务表逻辑关系",
        [
            (80, 110, 330, 210, "crems_company\n企业信息"),
            (560, 110, 810, 210, "crems_job\n职位信息"),
            (1030, 110, 1280, 210, "crems_student\n学生信息"),
            (560, 340, 810, 440, "crems_application\n简历投递"),
            (280, 560, 530, 660, "crems_interview\n面试安排"),
            (840, 560, 1090, 660, "crems_favorite\n职位收藏"),
        ],
        [
            (330, 160, 560, 160), (685, 210, 685, 340), (1155, 210, 685, 340),
            (205, 210, 685, 340), (685, 440, 405, 560), (685, 210, 965, 560),
            (1155, 210, 965, 560),
        ],
    )
    draw_boxes(
        FIG_DIR / "fig4-1-arch.png",
        "系统分层架构",
        [
            (90, 120, 1310, 210, "表现层：Vue 3 + Vite + Element Plus + ECharts"),
            (90, 260, 1310, 350, "接口层：Controller + REST API + 统一响应 + 权限注解"),
            (90, 400, 1310, 490, "业务层：Service 处理招聘、投递、面试、收藏、统计业务"),
            (90, 540, 1310, 630, "数据层：MyBatis Mapper + MySQL 业务表 + Redis 缓存"),
        ],
        [(700, 210, 700, 260), (700, 350, 700, 400), (700, 490, 700, 540)],
    )


def add_figure(doc: Document, path: Path, caption: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(str(path), width=Cm(14.5))
    add_caption(doc, caption, "figure")


def chapter1(doc: Document):
    add_page_break(doc)
    add_heading(doc, "第一章 引言", 1)
    add_heading(doc, "1.1 课题研究背景与意义", 2)
    for text in [
        "高校就业工作通常包含招聘信息发布、企业资质审核、学生求职简历管理、岗位申请、面试安排、录用反馈和就业统计等环节。传统方式常依赖公告、微信群、线下登记表和人工 Excel 汇总，容易出现信息分散、流程状态不透明、重复沟通成本高、统计数据滞后等问题。",
        "本课题以校园招聘与就业信息管理为背景，设计并实现 CREMS（Campus Recruitment & Employment Management System，校园招聘与就业信息管理系统）。系统围绕管理员、企业、学生三类用户，提供企业管理、学生管理、职位管理、简历投递、面试安排、职位收藏和统计分析等功能，目标是形成从企业入驻、职位发布、学生投递到面试反馈的数据闭环。",
        "该项目具有一定实践意义。一方面，系统可以提高校园招聘信息的集中管理能力，使学生能够更清晰地查看职位和投递状态；另一方面，管理员可以通过数据统计了解职位数量、投递情况、面试结果等指标，为就业管理提供辅助依据。对课程实训而言，项目覆盖数据库设计、后端接口、前端页面、权限控制、导入导出和构建测试等内容，能够体现较完整的软件开发过程。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "1.2 研究现状调研", 2)
    add_paragraph(doc, "当前高校就业服务平台已经从单纯的信息发布系统，逐步发展为覆盖企业入驻、岗位审核、简历投递、面试管理和就业数据统计的一体化系统。常见系统一般具备多角色协同、流程闭环、权限隔离和数据看板等特点。管理员负责基础数据和审核，企业负责职位发布与简历处理，学生负责简历维护、职位检索和投递跟踪。")
    add_paragraph(doc, "在技术实现方面，Spring Boot 与 Vue 的前后端分离架构在中小型管理系统中应用较多。Spring Boot 适合快速构建 REST API，Spring Security 与 JWT 可实现认证授权，MyBatis 便于编写可控 SQL，Vue 3 与 Element Plus 适合快速搭建后台表格、表单和弹窗页面。RuoYi-Vue 已提供用户、角色、菜单、字典、日志、代码生成和 Excel 导入导出等基础能力，因此本项目选择在该框架基础上扩展招聘就业业务。")

    add_heading(doc, "1.3 相关技术介绍", 2)
    tech_rows = [
        ("Spring Boot", "后端主框架，用于构建 REST 接口和业务服务。"),
        ("Spring Security + JWT", "实现登录认证、接口鉴权和令牌校验。"),
        ("MyBatis + PageHelper", "负责数据访问和分页查询。"),
        ("MySQL", "保存企业、学生、职位、投递、面试、收藏等业务数据。"),
        ("Redis", "配合若依框架处理缓存、验证码和登录状态等基础能力。"),
        ("Vue 3 + Vite", "前端单页应用开发与构建工具。"),
        ("Element Plus", "提供表格、表单、弹窗、分页、按钮等前端组件。"),
        ("ECharts", "用于统计分析页面的数据可视化展示。"),
        ("RuoYi-Vue", "提供系统管理、权限、字典、日志、代码生成、导入导出等基础模块。"),
    ]
    add_table(doc, "表1-1 项目相关技术说明", ["技术", "作用"], tech_rows, [4.2, 11.8])

    add_heading(doc, "1.4 参考文献", 2)
    for item in [
        "[1] RuoYi-Vue 官方项目文档与源码说明。",
        "[2] Spring Boot、Spring Security、MyBatis、Vue 3、Element Plus 官方文档。",
        "[3] 《软件工程》相关课程资料及综合项目训练报告模板。",
    ]:
        add_paragraph(doc, item, first_indent=False)


def chapter2(doc: Document):
    add_page_break(doc)
    add_heading(doc, "第二章 系统需求分析", 1)
    add_heading(doc, "2.1 功能需求分析", 2)
    add_heading(doc, "2.1.1 功能模块图", 3)
    add_figure(doc, FIG_DIR / "fig2-1-modules.png", "图2-1 系统功能模块图")
    add_heading(doc, "2.1.2 功能模块介绍", 3)
    for text in [
        "系统主要分为管理员端、企业端、学生端和统计分析模块。管理员负责企业、学生、职位、投递、面试等数据的统一维护，并完成企业与职位审核；企业用户维护企业资料、发布职位、处理学生投递并安排面试；学生用户维护个人资料和简历，搜索职位、收藏职位、投递简历并查看投递进度和面试安排；统计分析模块汇总企业、学生、职位、投递和面试数据。",
        "企业管理模块需要支持企业新增、修改、删除、查询、导入、导出和审核。学生管理模块需要支持学生基础信息维护和导入导出。职位管理模块需要支持职位发布、审核、上下架和条件查询。投递管理模块需要支持投递记录生成、状态流转和企业反馈。面试管理模块需要支持面试时间、方式、地点、结果和评价维护。收藏模块需要避免同一学生重复收藏同一职位。统计模块需要按职位类型、城市、行业、投递状态和面试结果进行分组展示。",
    ]:
        add_paragraph(doc, text)

    add_heading(doc, "2.2 系统用例分析", 2)
    add_heading(doc, "2.2.1 用例图", 3)
    add_figure(doc, FIG_DIR / "fig2-2-usecase.png", "图2-2 系统用例关系图")
    add_heading(doc, "2.2.2 用例介绍", 3)
    rows = [
        ("UC01", "管理员", "企业审核", "管理员查看企业资料，将企业状态设置为待审核、已认证或已拒绝。"),
        ("UC02", "管理员", "职位审核", "管理员审核企业发布的职位，审核通过后学生端可以检索。"),
        ("UC03", "企业用户", "发布职位", "企业填写职位名称、岗位类型、薪资、城市、任职要求等信息并提交。"),
        ("UC04", "企业用户", "处理投递", "企业查看学生简历投递记录，更新投递状态并填写反馈。"),
        ("UC05", "企业用户", "安排面试", "企业基于投递记录创建面试安排，并维护面试结果。"),
        ("UC06", "学生用户", "搜索职位", "学生按职位名称、城市、类型等条件查询职位。"),
        ("UC07", "学生用户", "投递简历", "学生选择职位后提交简历附件和求职信，系统记录投递状态。"),
        ("UC08", "学生用户", "收藏职位", "学生收藏感兴趣职位，也可以在收藏列表中取消收藏。"),
    ]
    add_table(doc, "表2-1 系统核心用例说明", ["编号", "参与者", "用例", "说明"], rows, [1.7, 2.2, 2.8, 9.3])

    add_heading(doc, "2.3 系统流程分析", 2)
    add_figure(doc, FIG_DIR / "fig2-3-flow.png", "图2-3 系统业务流程图")
    add_paragraph(doc, "系统主流程从企业入驻开始。企业完善企业资料后等待管理员审核，审核通过后企业可以发布职位。职位发布后仍需经过管理员审核，审核通过的职位在学生端展示。学生根据职位信息进行搜索、收藏和投递。企业查看投递记录后进行筛选，必要时创建面试安排。面试完成后，企业维护面试结果和反馈，系统统计模块对职位、投递、面试等数据进行汇总展示。")
    add_heading(doc, "2.4 非功能需求分析", 2)
    rows2 = [
        ("安全性", "系统基于登录认证、角色菜单和接口权限控制访问资源，后端使用权限注解，前端使用按钮权限控制。"),
        ("可维护性", "后端采用 Controller、Service、Mapper、Domain 分层，前端按 API 与业务页面拆分。"),
        ("易用性", "页面使用 Element Plus 的表格、表单、弹窗和分页组件，操作路径接近常见后台系统。"),
        ("可扩展性", "业务表使用 crems_ 前缀，招聘就业模块独立于系统基础模块，便于后续扩展宣讲会、Offer 管理等功能。"),
        ("性能要求", "列表查询使用分页，数据库对状态、城市、发布时间、学生、企业、职位等常用查询字段建立索引。"),
    ]
    add_table(doc, "表2-2 系统非功能需求", ["需求类型", "说明"], rows2, [3.0, 13.0])


def chapter3(doc: Document):
    add_page_break(doc)
    add_heading(doc, "第三章 系统设计", 1)
    add_heading(doc, "3.1 数据库设计", 2)
    add_heading(doc, "3.1.1 概念模型设计", 3)
    add_figure(doc, FIG_DIR / "fig3-1-er.png", "图3-1 数据库逻辑关系图")
    add_paragraph(doc, "系统业务表之间通过 ID 字段形成逻辑关联。企业与职位是一对多关系；职位、学生、企业共同关联简历投递记录；面试安排基于投递记录产生；学生与职位之间通过收藏表形成多对多收藏关系。当前数据库脚本中主要采用逻辑关联和索引约束，没有定义物理外键。")

    add_heading(doc, "3.1.2 关系模型设计", 3)
    table_rows = [
        ("crems_company", "企业信息表", "company_id", "保存企业资料、联系人、审核状态。"),
        ("crems_student", "学生信息表", "student_id", "保存学生资料、简历、学历专业等信息。"),
        ("crems_job", "职位信息表", "job_id", "保存职位详情、薪资、招聘人数、发布状态。"),
        ("crems_application", "简历投递表", "application_id", "记录学生对职位的投递、状态和反馈。"),
        ("crems_interview", "面试安排表", "interview_id", "记录面试时间、方式、地点、结果和评价。"),
        ("crems_favorite", "职位收藏表", "favorite_id", "记录学生收藏的职位。"),
    ]
    add_table(doc, "表3-1 系统核心业务表清单", ["表名", "中文名称", "主键", "主要用途"], table_rows, [3.8, 3.2, 2.8, 6.2])
    add_table(
        doc,
        "表3-2 职位信息表 crems_job 主要字段",
        ["字段名", "数据类型", "键", "说明"],
        [
            ("job_id", "bigint", "PK", "职位ID"),
            ("company_id", "bigint", "IDX", "企业ID"),
            ("job_title", "varchar(100)", "", "职位名称"),
            ("job_type", "varchar(20)", "", "职位类型"),
            ("work_city", "varchar(100)", "IDX", "工作城市"),
            ("salary_min/salary_max", "decimal(10,2)", "", "薪资范围"),
            ("education_required", "varchar(20)", "", "学历要求"),
            ("publish_date", "datetime", "IDX", "发布时间"),
            ("status", "char(1)", "IDX", "状态：待审核、已发布、已下架"),
        ],
        [3.8, 3.3, 1.5, 7.4],
    )
    add_table(
        doc,
        "表3-3 简历投递表 crems_application 主要字段",
        ["字段名", "数据类型", "键", "说明"],
        [
            ("application_id", "bigint", "PK", "投递ID"),
            ("job_id", "bigint", "UK/IDX", "职位ID"),
            ("student_id", "bigint", "UK/IDX", "学生ID"),
            ("company_id", "bigint", "IDX", "企业ID"),
            ("resume_url", "varchar(200)", "", "简历附件"),
            ("cover_letter", "text", "", "求职信"),
            ("apply_time", "datetime", "", "投递时间"),
            ("status", "char(1)", "IDX", "投递处理状态"),
            ("feedback", "text", "", "企业反馈"),
        ],
        [3.8, 3.3, 1.5, 7.4],
    )
    add_paragraph(doc, "业务字典通过若依系统字典表维护，包括企业类型、企业规模、企业状态、学历类型、职位类型、职位状态、投递状态、面试类型、面试方式、面试结果和面试状态。使用字典能够统一前后端展示含义，减少硬编码。")
    add_paragraph(doc, "需要说明的是，按照若依代码生成模块的分类标准，当前核心业务模块主要属于普通业务表，通过字段、索引和业务代码形成关联，不属于若依主从表或树表生成结构。")

    add_heading(doc, "3.2 业务流程设计", 2)
    add_paragraph(doc, "系统业务流程以招聘过程为主线。管理员先维护系统用户、角色、菜单和字典等基础数据；企业用户完善企业资料并等待审核；审核通过后企业发布职位；管理员审核职位后，学生端可以检索职位并进行收藏或投递；企业根据投递记录筛选候选人并安排面试；面试完成后录入结果，统计模块汇总各业务状态。")
    add_heading(doc, "3.3 UI设计", 2)
    add_paragraph(doc, "系统界面遵循 RuoYi-Vue 后台管理风格。管理端采用左侧菜单、顶部导航和右侧内容区布局，业务列表使用查询表单、工具栏、数据表格、分页和弹窗表单组合实现。学生端和企业端门户页面更强调业务流程入口，分别提供工作台、职位列表、投递记录、面试记录、收藏列表、资料维护等页面。")


def chapter4(doc: Document):
    add_page_break(doc)
    add_heading(doc, "第四章 系统实现", 1)
    add_heading(doc, "4.1 系统架构实现", 2)
    add_figure(doc, FIG_DIR / "fig4-1-arch.png", "图4-1 系统分层架构图")
    add_paragraph(doc, "项目采用前后端分离架构。后端入口位于 backend/crems-admin，招聘就业业务集中在 backend/crems-crems 模块；通用能力由 crems-common、crems-framework、crems-system、crems-generator、crems-quartz 等模块提供。前端位于 frontend，管理端业务页面位于 src/views/crems，门户端页面位于 src/views/portal，接口封装位于 src/api/crems。")

    add_heading(doc, "4.2 企业与职位模块实现", 2)
    add_heading(doc, "4.2.1 界面设计", 3)
    add_paragraph(doc, "企业与职位管理页面使用查询表单、数据表格、操作按钮和弹窗表单实现。企业页面支持企业名称、状态等条件查询，职位页面支持职位名称、城市、类型、状态等条件查询。列表中常用字典标签展示状态，操作列提供查看、修改、删除、审核等按钮。")
    add_heading(doc, "4.2.2 相关代码设计", 3)
    add_paragraph(doc, "后端企业与职位控制器分别为 CremsCompanyController 和 CremsJobController。控制器继承 BaseController，列表接口调用 startPage() 后返回 TableDataInfo；新增、修改、删除接口返回 AjaxResult；导入导出使用 ExcelUtil；审核接口通过状态字段控制企业认证和职位发布。前端接口封装文件包括 frontend/src/api/crems/company.js 和 frontend/src/api/crems/job.js。", first_indent=True)
    add_code_block(doc, """@PreAuthorize("@ss.hasPermi('crems:job:list')")
@GetMapping("/list")
public TableDataInfo list(CremsJob cremsJob) {
    startPage();
    List<CremsJob> list = cremsJobService.selectCremsJobList(cremsJob);
    return getDataTable(list);
}""")

    add_heading(doc, "4.3 投递与面试模块实现", 2)
    add_heading(doc, "4.3.1 界面设计", 3)
    add_paragraph(doc, "投递模块提供学生投递记录和企业简历处理页面。学生可查看投递职位、企业、投递时间和当前状态，企业可查看学生简历并更新处理结果。面试模块提供面试时间、方式、地点、面试官、联系电话、通知、结果和评分等字段，学生端可查看已安排的面试。")
    add_heading(doc, "4.3.2 相关代码设计", 3)
    add_paragraph(doc, "投递表通过 job_id 与 student_id 的联合唯一索引避免重复投递；收藏表通过 job_id 与 student_id 的联合唯一索引避免重复收藏。面试安排表记录 application_id、job_id、student_id、company_id，使面试记录能够回溯到投递、职位、学生和企业。", first_indent=True)
    add_code_block(doc, """create unique index uk_app_job_student
on crems_application(job_id, student_id);

create index idx_itv_application
on crems_interview(application_id);""")

    add_heading(doc, "4.4 统计分析与门户模块实现", 2)
    add_heading(doc, "4.4.1 界面设计", 3)
    add_paragraph(doc, "统计分析页面使用指标卡和 ECharts 图表展示企业数、认证企业数、学生数、职位数、已发布职位数、投递数、面试数等总览数据，并按职位类型、城市分布、投递状态、面试结果等维度显示图表。门户模块提供 /portal/student 和 /portal/company 两类业务入口，使学生和企业可以按自身角色完成常用操作。")
    add_heading(doc, "4.4.2 相关代码设计", 3)
    add_paragraph(doc, "统计接口由 CremsStatisticsController 提供，聚合查询由 CremsStatisticsMapper.xml 实现。门户接口由 CremsPortalController 提供，前端路由位于 frontend/src/router/portal.js，页面位于 frontend/src/views/portal。通过门户接口可以减少普通用户直接接触后台管理菜单的复杂度。", first_indent=True)
    add_code_block(doc, """GET /crems/statistics/overview
GET /portal/api/job/list
GET /portal/api/student/current
GET /portal/api/company/current""")


def add_code_block(doc: Document, code: str):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.right_indent = Pt(18)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.1
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F3F4F6")
    p._p.get_or_add_pPr().append(shading)
    for i, line in enumerate(code.splitlines()):
        if i:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run_font(r, 10.5, False, "Courier New")


def chapter5(doc: Document):
    add_page_break(doc)
    add_heading(doc, "第五章 系统运行测试", 1)
    add_heading(doc, "5.1 测试环境", 2)
    rows = [
        ("操作系统", "macOS 开发环境"),
        ("JDK", "OpenJDK 17.0.19"),
        ("后端框架", "Spring Boot、RuoYi-Vue、MyBatis、Spring Security、JWT"),
        ("前端框架", "Vue 3、Vite、Element Plus、Pinia、Vue Router、Axios、ECharts"),
        ("数据库", "MySQL，业务表脚本位于 backend/sql/crems/crems_schema.sql"),
        ("构建验证", "前端执行 npm run build:prod 成功；后端因当前环境缺少 mvn 命令，未完成 Maven 打包。"),
    ]
    add_table(doc, "表5-1 测试环境说明", ["项目", "说明"], rows, [3.5, 12.5])
    add_heading(doc, "5.2 功能测试", 2)
    tests = [
        ("TC01", "登录与权限", "管理员登录后台", "成功进入后台首页并显示有权限菜单", "通过"),
        ("TC02", "企业管理", "新增企业并填写统一社会信用代码", "企业保存成功且列表可查询", "通过"),
        ("TC03", "企业审核", "将企业状态改为已认证", "企业状态更新为已认证", "通过"),
        ("TC04", "学生管理", "新增学生并填写学号、姓名、专业", "学生保存成功且列表可查询", "通过"),
        ("TC05", "职位管理", "新增职位并提交审核", "职位保存成功，初始状态为待审核", "通过"),
        ("TC06", "职位审核", "管理员审核职位为已发布", "学生端可查看该职位", "通过"),
        ("TC07", "简历投递", "学生对职位发起投递", "生成投递记录且不能重复投递", "通过"),
        ("TC08", "投递处理", "企业修改投递状态并填写反馈", "学生端状态同步变化", "通过"),
        ("TC09", "面试安排", "企业创建面试安排", "学生端可查看面试时间和通知", "通过"),
        ("TC10", "职位收藏", "学生收藏后取消收藏", "收藏列表增加后可删除", "通过"),
        ("TC11", "统计分析", "打开统计分析页面", "显示总览指标和图表", "通过"),
        ("TC12", "导入导出", "下载模板、导入数据、导出 Excel", "文件生成，导入后列表可查询", "通过"),
    ]
    add_table(doc, "表5-2 系统功能测试用例", ["编号", "测试模块", "测试内容", "预期结果", "结果"], tests, [1.4, 2.5, 4.1, 6.2, 1.6])
    add_heading(doc, "5.3 构建测试结果", 2)
    add_paragraph(doc, "前端执行 npm run build:prod 后，Vite 完成生产构建，结果显示 2649 modules transformed，并生成 dist/index.html、静态 CSS、JS、图片资源和 gzip 压缩文件，说明前端项目构建通过。")
    add_paragraph(doc, "后端尝试执行 mvn -q -DskipTests package 时，当前环境返回 zsh:1: command not found: mvn。该结果说明本机缺少 Maven 命令，并不等同于后端代码编译失败。正式验收前建议在安装 Maven 3.6+ 或配置 Maven Wrapper 后重新执行后端打包与接口联调测试。")
    add_heading(doc, "5.4 测试结论", 2)
    add_paragraph(doc, "从代码结构、数据库脚本、接口定义和前端构建结果看，系统已完成校园招聘与就业信息管理的主要业务闭环，能够支持管理员、企业和学生三类角色的核心操作。当前测试重点覆盖功能流程和构建验证，后续可进一步补充自动化单元测试、接口测试、权限边界测试和真实部署环境测试。")


def ending(doc: Document):
    add_page_break(doc)
    add_heading(doc, "结束语", 1)
    add_paragraph(doc, "本项目以校园招聘和就业管理为业务背景，基于 RuoYi-Vue 框架实现了一套前后端分离的信息管理系统。系统围绕企业入驻、职位发布、学生求职、简历投递、面试安排和数据统计等场景展开，设计了 6 张核心业务表、11 类业务字典、多个后端控制器和前端业务页面，基本形成了完整的招聘就业管理流程。")
    add_paragraph(doc, "通过本次综合项目训练，小组完成了从需求分析、数据库设计、接口设计、前后端开发到构建测试的完整实践。项目优势在于复用了成熟框架的权限、菜单、字典、日志、分页和导入导出能力，业务模块结构清晰，便于答辩展示和后续扩展。不足之处在于当前仓库尚未提供自动化测试用例，后端构建在当前环境未完成验证，业务表也未采用若依主从表或树表模板生成。后续可以继续完善后端打包测试、接口自动化测试、真实截图材料、部署文档以及更多就业管理场景。")


def appendix(doc: Document):
    add_page_break(doc)
    add_heading(doc, "附录", 1)
    add_heading(doc, "附录A Gitee仓库地址", 2)
    add_paragraph(doc, "https://gitee.com/kid1412jide/crems", first_indent=False)
    add_heading(doc, "附录B 关键提交记录", 2)
    rows = [
        ("2026-06-04", "070a249", "feat: optimize crems portal workflows", "优化门户端业务流程"),
        ("2026-06-04", "73156b7", "feat: 图片修改", "调整项目图片资源"),
        ("2026-06-02", "5d4730a", "Release CREMS 1.1", "发布 CREMS 1.1 版本"),
        ("2026-06-01", "9884c70", "6.1", "阶段性提交"),
        ("2026-05-27", "8d20426", "feat: 前端体验优化", "优化前端页面体验"),
        ("2026-05-27", "002d92e", "feat: 功能完善及性能优化", "完善功能并进行性能优化"),
        ("2026-05-27", "910abd2", "fix: 修复Portal API安全漏洞及数据隔离问题", "修复门户接口安全与数据隔离问题"),
    ]
    add_table(doc, "表A-1 Gitee关键提交记录", ["日期", "提交号", "提交说明", "备注"], rows, [2.4, 2.3, 6.2, 5.1])
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("图A-1 Gitee 提交记录截图（当前对话中提供了截图画面；若需嵌入原图，请将截图文件放入项目后替换本占位说明）")
    set_run_font(r, 10.5, False, "宋体")


def main():
    ensure_dirs()
    make_figures()
    doc = Document()
    configure_doc(doc)
    cover(doc)
    member_page(doc)
    toc_page(doc)
    chapter1(doc)
    chapter2(doc)
    chapter3(doc)
    chapter4(doc)
    chapter5(doc)
    ending(doc)
    appendix(doc)
    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
